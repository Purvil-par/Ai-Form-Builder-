"""
File Parser Utility Module
Parses PDF, DOCX, XLSX files and extracts text content for LLM processing.

Instead of sending Base64 encoded binary to LLM (which LLM cannot understand),
this module extracts actual text content that LLM can process.

Supports OCR for image-based/scanned PDFs using Tesseract OCR.
"""

import base64
import io
import logging
import re
import os
from typing import Optional, Dict, Any

# PDF parsing
import fitz  # PyMuPDF

# DOCX parsing
from docx import Document

# XLSX parsing
from openpyxl import load_workbook

logger = logging.getLogger("file-parser")

# Maximum characters to send to LLM (to prevent token overflow)
MAX_CONTENT_LENGTH = 50000

# Minimum text threshold - if page has less text than this, use OCR
MIN_TEXT_THRESHOLD = 50

# ============================================
# OCR Setup with Tesseract
# ============================================
OCR_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    
    # Auto-detect Tesseract on Windows
    TESSERACT_PATHS = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        r"C:\Users\Purvil\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
        os.environ.get("TESSERACT_CMD", ""),
    ]
    
    tesseract_found = False
    for path in TESSERACT_PATHS:
        if path and os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            tesseract_found = True
            logger.info(f"✅ Tesseract OCR found at: {path}")
            break
    
    if not tesseract_found:
        # Try default path (might be in PATH)
        try:
            pytesseract.get_tesseract_version()
            tesseract_found = True
            logger.info("✅ Tesseract OCR found in system PATH")
        except:
            pass
    
    if tesseract_found:
        OCR_AVAILABLE = True
        logger.info("✅ OCR support ENABLED - image-based PDFs will be processed")
    else:
        logger.warning("⚠️ Tesseract not found - OCR disabled. Install from: https://github.com/UB-Mannheim/tesseract/wiki")
        
except ImportError as e:
    logger.warning(f"⚠️ OCR libraries not installed: {e}")
    logger.warning("   Install with: pip install pytesseract Pillow")


def extract_text_with_ocr(page) -> str:
    """
    Extract text from a PDF page using Tesseract OCR.
    
    Args:
        page: PyMuPDF page object
        
    Returns:
        Extracted text from the page image
    """
    if not OCR_AVAILABLE:
        return ""
    
    try:
        # Render page at 2x resolution for better OCR accuracy
        zoom = 2.0
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        
        # Convert to PIL Image
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        
        # Run OCR with English language
        text = pytesseract.image_to_string(img, lang='eng')
        
        return text.strip()
        
    except Exception as e:
        logger.warning(f"⚠️ OCR failed for page: {e}")
        return ""


# ============================================
# Image Analysis for Form Generation
# ============================================

def extract_text_from_image(base64_image: str) -> Dict[str, Any]:
    """
    Extract text from a Base64 encoded image using OCR.
    
    Args:
        base64_image: Base64 encoded image (can include data URL prefix)
        
    Returns:
        Dictionary with extracted text and metadata
    """
    if not OCR_AVAILABLE:
        logger.warning("⚠️ OCR not available for image text extraction")
        return {
            "success": False,
            "error": "OCR not available",
            "text": ""
        }
    
    try:
        # Remove data URL prefix if present
        if "," in base64_image:
            base64_image = base64_image.split(",")[1]
        
        # Decode Base64 to bytes
        image_bytes = base64.b64decode(base64_image)
        
        # Open image with PIL
        img = Image.open(io.BytesIO(image_bytes))
        
        # Convert to RGB if necessary (for PNG with transparency)
        if img.mode in ('RGBA', 'P'):
            img = img.convert('RGB')
        
        # Run OCR
        extracted_text = pytesseract.image_to_string(img, lang='eng')
        
        # Get image info
        width, height = img.size
        
        result = {
            "success": True,
            "text": extracted_text.strip(),
            "char_count": len(extracted_text.strip()),
            "image_width": width,
            "image_height": height,
            "has_text": len(extracted_text.strip()) > 20
        }
        
        logger.info(f"✅ Image OCR: {result['char_count']} chars extracted from {width}x{height} image")
        
        return result
        
    except Exception as e:
        logger.error(f"❌ Image OCR failed: {e}")
        return {
            "success": False,
            "error": str(e),
            "text": ""
        }


def analyze_image_for_form_generation(base64_image: str) -> Dict[str, Any]:
    """
    Analyze an image and prepare it for form generation.
    
    This function:
    1. Extracts text using OCR (for text-heavy images)
    2. Prepares metadata for Vision API analysis
    3. Classifies content type based on OCR results
    
    Args:
        base64_image: Base64 encoded image (with or without data URL prefix)
        
    Returns:
        Dictionary with analysis results for LLM processing
    """
    logger.info("📷 Analyzing image for form generation...")
    
    # Clean base64 string
    clean_base64 = base64_image
    image_mime_type = "image/jpeg"
    
    if "," in base64_image:
        # Extract mime type from data URL
        header = base64_image.split(",")[0]
        if "image/png" in header:
            image_mime_type = "image/png"
        elif "image/gif" in header:
            image_mime_type = "image/gif"
        elif "image/webp" in header:
            image_mime_type = "image/webp"
        clean_base64 = base64_image.split(",")[1]
    
    # Extract text using OCR
    ocr_result = extract_text_from_image(base64_image)
    
    # Classify content based on OCR results
    content_type = "visual"  # Default: diagram, chart, or nature image
    extracted_text = ocr_result.get("text", "")
    
    if ocr_result.get("has_text") and len(extracted_text) > 100:
        # Check for code patterns
        code_indicators = ["def ", "function ", "class ", "import ", "const ", "let ", "var ", "public ", "private ", "return ", "if (", "for (", "while (", "=>", "->"]
        if any(indicator in extracted_text for indicator in code_indicators):
            content_type = "code"
        else:
            content_type = "text_document"
    elif ocr_result.get("has_text"):
        content_type = "mixed"  # Some text with visual elements
    
    result = {
        "success": True,
        "content_type": content_type,
        "mime_type": image_mime_type,
        "base64_data": clean_base64,
        "ocr_text": extracted_text,
        "ocr_char_count": len(extracted_text),
        "has_extractable_text": ocr_result.get("has_text", False),
        "image_dimensions": {
            "width": ocr_result.get("image_width", 0),
            "height": ocr_result.get("image_height", 0)
        }
    }
    
    logger.info(f"✅ Image analysis complete: type={content_type}, OCR chars={len(extracted_text)}")
    
    return result


def format_image_analysis_for_llm(analysis: Dict[str, Any], user_prompt: str) -> str:
    """
    Format image analysis results into a prompt for the LLM.
    
    Args:
        analysis: Result from analyze_image_for_form_generation()
        user_prompt: User's instructions for form generation
        
    Returns:
        Formatted prompt string to send to Vision LLM
    """
    content_type_descriptions = {
        "text_document": "text-heavy document (notes, slides, educational content)",
        "code": "code or programming content",
        "mixed": "mixed content (text with diagrams/images)",
        "visual": "visual content (diagram, chart, infographic, or image)"
    }
    
    content_desc = content_type_descriptions.get(analysis["content_type"], "unknown content")
    
    prompt_parts = [
        f"USER INSTRUCTIONS: {user_prompt}",
        "",
        f"IMAGE ANALYSIS:",
        f"- Content Type: {content_desc}",
        f"- Image Size: {analysis['image_dimensions']['width']}x{analysis['image_dimensions']['height']} pixels",
    ]
    
    # Add OCR text if available
    if analysis.get("has_extractable_text") and analysis.get("ocr_text"):
        prompt_parts.extend([
            "",
            "EXTRACTED TEXT FROM IMAGE (via OCR):",
            "---",
            analysis["ocr_text"][:5000],  # Limit to 5000 chars
            "---" if len(analysis["ocr_text"]) <= 5000 else "--- [truncated] ---",
        ])
    
    prompt_parts.extend([
        "",
        "TASK:",
        "1. Analyze the uploaded image carefully",
        "2. Identify key concepts, facts, or information in the image",
        "3. Follow the user's instructions to generate the form",
        "4. Create questions that test understanding of the image content",
        "5. Use appropriate question types (MCQ, short answer, etc.) based on the content",
        "",
        "⚠️ STRICT EXTRACTION RULES:",
        "- Generate form fields ONLY from content EXPLICITLY visible in the image",
        "- Do NOT add inferred or assumed questions beyond what is shown",
        "- Do NOT create duplicate or semantically similar questions",
        "- If the image has 10 questions, generate EXACTLY 10 fields - no more, no less",
        "- Preserve EXACT wording of any text/questions visible in the image",
        "- NEVER add 'bonus' or 'extra' questions not present in the source",
    ])
    
    return "\n".join(prompt_parts)


def parse_pdf_from_base64(base64_content: str) -> Dict[str, Any]:
    """
    Extract text content from a Base64 encoded PDF file.
    Uses OCR as fallback for image-based/scanned pages.
    
    Args:
        base64_content: Base64 encoded PDF data
        
    Returns:
        Dictionary with extracted content and metadata
    """
    try:
        # Decode Base64 to bytes
        pdf_bytes = base64.b64decode(base64_content)
        
        # Open PDF from bytes
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        # Extract text from all pages
        text_content = []
        ocr_pages = 0
        text_pages = 0
        
        for page_num, page in enumerate(doc, 1):
            # First try regular text extraction
            page_text = page.get_text("text").strip()
            extraction_method = "text"
            
            # If text is too short, try OCR
            if len(page_text) < MIN_TEXT_THRESHOLD:
                if OCR_AVAILABLE:
                    logger.info(f"📷 Page {page_num}: Only {len(page_text)} chars, using OCR...")
                    ocr_text = extract_text_with_ocr(page)
                    
                    if ocr_text and len(ocr_text) > len(page_text):
                        page_text = ocr_text
                        ocr_pages += 1
                        extraction_method = "OCR"
                        logger.info(f"✅ Page {page_num}: OCR extracted {len(ocr_text)} chars")
                    else:
                        logger.info(f"📄 Page {page_num}: OCR didn't improve, using original text")
                else:
                    logger.warning(f"⚠️ Page {page_num}: Low text ({len(page_text)} chars) but OCR not available")
            else:
                text_pages += 1
            
            if page_text:
                text_content.append(f"--- Page {page_num} ({extraction_method}) ---\n{page_text}")
        
        full_text = "\n\n".join(text_content)
        
        result = {
            "success": True,
            "file_type": "pdf",
            "page_count": len(doc),
            "text_pages": text_pages,
            "ocr_pages": ocr_pages,
            "char_count": len(full_text),
            "content": full_text,
            "truncated": False
        }
        
        doc.close()
        
        # Truncate if too large
        if len(full_text) > MAX_CONTENT_LENGTH:
            result["content"] = full_text[:MAX_CONTENT_LENGTH] + "\n\n[... Content truncated due to length ...]"
            result["truncated"] = True
        
        # Log summary
        if ocr_pages > 0:
            logger.info(f"✅ PDF parsed: {len(doc)} pages ({text_pages} text, {ocr_pages} OCR), {len(full_text)} chars")
        else:
            logger.info(f"✅ PDF parsed: {result['page_count']} pages, {result['char_count']} chars, truncated={result['truncated']}")
        
        return result
        
    except Exception as e:
        logger.error(f"❌ PDF parsing failed: {e}")
        return {
            "success": False,
            "error": str(e),
            "file_type": "pdf"
        }


def parse_docx_from_base64(base64_content: str) -> Dict[str, Any]:
    """
    Extract text content from a Base64 encoded DOCX file.
    
    Args:
        base64_content: Base64 encoded DOCX data
        
    Returns:
        Dictionary with extracted content and metadata
    """
    try:
        # Decode Base64 to bytes
        docx_bytes = base64.b64decode(base64_content)
        
        # Open DOCX from bytes
        doc = Document(io.BytesIO(docx_bytes))
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract text from tables
        tables_text = []
        for table_idx, table in enumerate(doc.tables, 1):
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(row_cells))
            if table_rows:
                tables_text.append(f"--- Table {table_idx} ---\n" + "\n".join(table_rows))
        
        # Combine content
        full_text = "\n\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n" + "\n\n".join(tables_text)
        
        result = {
            "success": True,
            "file_type": "docx",
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "char_count": len(full_text),
            "content": full_text,
            "truncated": False
        }
        
        # Truncate if too large
        if len(full_text) > MAX_CONTENT_LENGTH:
            result["content"] = full_text[:MAX_CONTENT_LENGTH] + "\n\n[... Content truncated due to length ...]"
            result["truncated"] = True
        
        logger.info(f"✅ DOCX parsed: {result['paragraph_count']} paragraphs, {result['table_count']} tables, {result['char_count']} chars")
        
        return result
        
    except Exception as e:
        logger.error(f"❌ DOCX parsing failed: {e}")
        return {
            "success": False,
            "error": str(e),
            "file_type": "docx"
        }


def parse_xlsx_from_base64(base64_content: str) -> Dict[str, Any]:
    """
    Extract text content from a Base64 encoded XLSX file.
    
    Args:
        base64_content: Base64 encoded XLSX data
        
    Returns:
        Dictionary with extracted content and metadata
    """
    try:
        # Decode Base64 to bytes
        xlsx_bytes = base64.b64decode(base64_content)
        
        # Open XLSX from bytes
        wb = load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
        
        # Extract data from all sheets
        sheets_content = []
        total_rows = 0
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows = []
            
            for row in sheet.iter_rows(values_only=True):
                # Filter out empty rows
                if any(cell is not None for cell in row):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    rows.append(row_text)
                    total_rows += 1
            
            if rows:
                sheets_content.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))
        
        full_text = "\n\n".join(sheets_content)
        
        result = {
            "success": True,
            "file_type": "xlsx",
            "sheet_count": len(wb.sheetnames),
            "total_rows": total_rows,
            "char_count": len(full_text),
            "content": full_text,
            "truncated": False
        }
        
        wb.close()
        
        # Truncate if too large
        if len(full_text) > MAX_CONTENT_LENGTH:
            result["content"] = full_text[:MAX_CONTENT_LENGTH] + "\n\n[... Content truncated due to length ...]"
            result["truncated"] = True
        
        logger.info(f"✅ XLSX parsed: {result['sheet_count']} sheets, {result['total_rows']} rows, {result['char_count']} chars")
        
        return result
        
    except Exception as e:
        logger.error(f"❌ XLSX parsing failed: {e}")
        return {
            "success": False,
            "error": str(e),
            "file_type": "xlsx"
        }


def detect_and_parse_file_content(file_content: str) -> Optional[str]:
    """
    Detect file type from content wrapper and parse accordingly.
    
    Frontend sends binary files in this format:
    [BINARY FILE: filename.pdf]
    [TYPE: application/pdf]
    [SIZE: 500.00KB]
    [BASE64_CONTENT_START]
    ...base64 data...
    [BASE64_CONTENT_END]
    
    Text files are sent as plain text (no parsing needed).
    
    Args:
        file_content: Raw file content from frontend
        
    Returns:
        Parsed text content or original content if not binary
    """
    if not file_content:
        return None
    
    # Check if it's a binary file wrapper
    if "[BINARY FILE:" not in file_content:
        # Plain text file - return as-is
        logger.info("📄 Text file detected - using content directly")
        return file_content
    
    logger.info("📦 Binary file wrapper detected - parsing...")
    
    # Extract file info
    file_name_match = re.search(r'\[BINARY FILE: (.+?)\]', file_content)
    file_type_match = re.search(r'\[TYPE: (.+?)\]', file_content)
    base64_match = re.search(r'\[BASE64_CONTENT_START\]\n?(.*?)\n?\[BASE64_CONTENT_END\]', file_content, re.DOTALL)
    
    if not base64_match:
        logger.warning("⚠️ Could not extract Base64 content from wrapper")
        return file_content
    
    file_name = file_name_match.group(1) if file_name_match else "unknown"
    file_type = file_type_match.group(1) if file_type_match else ""
    base64_content = base64_match.group(1).strip()
    
    logger.info(f"📁 Parsing file: {file_name} (type: {file_type})")
    
    # Determine parser based on file type or extension
    result = None
    
    if "pdf" in file_type.lower() or file_name.lower().endswith(".pdf"):
        result = parse_pdf_from_base64(base64_content)
    elif "word" in file_type.lower() or file_name.lower().endswith((".docx", ".doc")):
        result = parse_docx_from_base64(base64_content)
    elif "sheet" in file_type.lower() or "excel" in file_type.lower() or file_name.lower().endswith((".xlsx", ".xls")):
        result = parse_xlsx_from_base64(base64_content)
    else:
        logger.warning(f"⚠️ Unsupported binary file type: {file_type}")
        return f"[Unsupported binary file: {file_name}]"
    
    if result and result.get("success"):
        # Format parsed content for LLM
        parsed_output = f"""[PARSED FILE: {file_name}]
[TYPE: {result['file_type'].upper()}]
[STATS: {result.get('page_count', result.get('paragraph_count', result.get('sheet_count', 'N/A')))} items, {result['char_count']} characters]
[TRUNCATED: {result['truncated']}]

--- EXTRACTED CONTENT ---
{result['content']}
--- END OF CONTENT ---"""
        
        # Log the parsed content for debugging (first 500 chars)
        logger.info(f"📝 PARSED CONTENT PREVIEW:\n{result['content'][:500]}...")
        
        return parsed_output
    else:
        error_msg = result.get("error", "Unknown parsing error") if result else "Parser returned None"
        logger.error(f"❌ File parsing failed: {error_msg}")
        return f"[Failed to parse file: {file_name}. Error: {error_msg}]"
